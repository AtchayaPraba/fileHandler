import os, glob, shutil
from docx import Document
import re

class fileHandler:
    
    def __init__ (self):
        self.cwd = os.getcwd()
        
    def readText (self, filename, filepath):
        os.chdir(filepath)
        try:
            with open(filename) as file:
                print (file.read())
        except Exception as e:
            print (e)
        else:
            print (filename, "READ SUCCESSFULLY !!")
        finally:
            os.chdir(self.cwd)
    
    def readDocument (self, filename, filepath):
        os.chdir(filepath)
        try:
            para_list = []
            file = Document(filepath+"\\"+filename)
            for para in file.paragraphs:
                para_list.append(para.text)
            print ("\n".join(para_list))
        except Exception as e:
            print (e)
        else:
            print (filename, "READ SUCCESSFULLY !!")
        finally:
            os.chdir(self.cwd)
            
    def appendDoctoText (self, filename, filepath, op_filename, op_filepath):
        os.chdir(filepath)
        try:
            para_list = []
            file = Document(filepath+"\\"+filename)
            for para in file.paragraphs:
                para_list.append(para.text)
            os.chdir(op_filepath)
            with open(op_filename, "a+", encoding="utf-8") as file:
                file.write(filename.upper().center(130, "*"))
                file.write("\n"*2)
                file.write("\n".join(para_list))
                file.write("\n"*2)
        except Exception as e:
            print (e)
        else:
            print (op_filename, "APPENDED SUCCESSFULLY !!")
        finally:
            os.chdir(self.cwd)
            
    def get_numbers_from_string (self,string):
        try:
            temp = re.findall(r'\d+', string)
            res = list(map(int, temp))
            return(res)
        except Exception as e:
            print (e)
        else:
            pass
        finally:
            pass
        
    def copyFile (self, src, dst, filename, new_filename): 
        try:
            shutil.copy(src=src+"\\"+filename, dst=dst+"\\"+new_filename)
        except Exception as e:
            print (e)
        else:
            print (filename, "IS COPPIED AS", new_filename)
        finally:
            os.chdir(self.cwd)